import io
from pathlib import Path

# Minimum characters per page to consider pdfplumber successful.
# Below this threshold we assume the PDF is scanned/image-based.
OCR_FALLBACK_THRESHOLD = 50


def extract_text(file_bytes: bytes, filename: str) -> dict:
    ext = Path(filename).suffix.lower()

    if ext == ".pdf":
        return _extract_pdf(file_bytes, filename)
    elif ext in (".docx", ".doc"):
        return _extract_docx(file_bytes, filename)
    elif ext == ".txt":
        return _extract_txt(file_bytes, filename)
    elif ext in (".png", ".jpg", ".jpeg", ".tiff", ".tif", ".bmp"):
        return _extract_image_ocr(file_bytes, filename)
    else:
        return {
            "filename": filename,
            "extension": ext,
            "pages": [],
            "full_text": "",
            "ocr_used": False,
            "error": f"Unsupported file type: '{ext}'. Supported: PDF, DOCX, TXT, PNG, JPG",
        }


def _extract_pdf(file_bytes: bytes, filename: str) -> dict:
    result = _extract_pdf_text(file_bytes, filename)

    # Check if we got meaningful text
    avg_chars = (
        sum(len(p["text"]) for p in result["pages"]) / max(len(result["pages"]), 1)
        if result["pages"] else 0
    )

    if not result["error"] and avg_chars < OCR_FALLBACK_THRESHOLD:
        # Scanned PDF — try OCR
        ocr_result = _extract_pdf_ocr(file_bytes, filename)
        if not ocr_result["error"] and ocr_result["full_text"].strip():
            return ocr_result

    return result


def _extract_pdf_text(file_bytes: bytes, filename: str) -> dict:
    try:
        import pdfplumber

        pages = []
        full_text_parts = []

        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for i, page in enumerate(pdf.pages, start=1):
                text = page.extract_text() or ""
                tables = page.extract_tables()
                table_text = _tables_to_text(tables)
                page_text = text + ("\n" + table_text if table_text else "")
                pages.append({"page": i, "text": page_text.strip()})
                full_text_parts.append(page_text.strip())

        return {
            "filename": filename,
            "extension": ".pdf",
            "pages": pages,
            "full_text": "\n\n".join(full_text_parts),
            "ocr_used": False,
            "error": None,
        }
    except Exception as e:
        return {
            "filename": filename,
            "extension": ".pdf",
            "pages": [],
            "full_text": "",
            "ocr_used": False,
            "error": str(e),
        }


def _extract_pdf_ocr(file_bytes: bytes, filename: str) -> dict:
    try:
        from pdf2image import convert_from_bytes
        import numpy as np

        reader = _get_ocr_reader()
        images = convert_from_bytes(file_bytes, dpi=200)

        pages = []
        full_text_parts = []

        for i, img in enumerate(images, start=1):
            img_array = np.array(img)
            results = reader.readtext(img_array, detail=0, paragraph=True)
            page_text = "\n".join(results)
            pages.append({"page": i, "text": page_text.strip()})
            full_text_parts.append(page_text.strip())

        return {
            "filename": filename,
            "extension": ".pdf",
            "pages": pages,
            "full_text": "\n\n".join(full_text_parts),
            "ocr_used": True,
            "error": None,
        }
    except Exception as e:
        return {
            "filename": filename,
            "extension": ".pdf",
            "pages": [],
            "full_text": "",
            "ocr_used": True,
            "error": f"OCR failed: {str(e)}",
        }


def _extract_image_ocr(file_bytes: bytes, filename: str) -> dict:
    try:
        import numpy as np
        from PIL import Image

        reader = _get_ocr_reader()
        img = Image.open(io.BytesIO(file_bytes)).convert("RGB")
        img_array = np.array(img)
        results = reader.readtext(img_array, detail=0, paragraph=True)
        text = "\n".join(results)

        return {
            "filename": filename,
            "extension": Path(filename).suffix.lower(),
            "pages": [{"page": 1, "text": text.strip()}],
            "full_text": text,
            "ocr_used": True,
            "error": None,
        }
    except Exception as e:
        return {
            "filename": filename,
            "extension": Path(filename).suffix.lower(),
            "pages": [],
            "full_text": "",
            "ocr_used": True,
            "error": f"OCR failed: {str(e)}",
        }



def _extract_docx(file_bytes: bytes, filename: str) -> dict:
    try:
        from docx import Document

        doc = Document(io.BytesIO(file_bytes))
        paragraphs = []
        full_parts = []

        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if text:
                paragraphs.append({"page": i + 1, "text": text})
                full_parts.append(text)

        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    full_parts.append(row_text)

        return {
            "filename": filename,
            "extension": ".docx",
            "pages": paragraphs,
            "full_text": "\n".join(full_parts),
            "ocr_used": False,
            "error": None,
        }
    except Exception as e:
        return {
            "filename": filename,
            "extension": ".docx",
            "pages": [],
            "full_text": "",
            "ocr_used": False,
            "error": str(e),
        }



def _extract_txt(file_bytes: bytes, filename: str) -> dict:
    try:
        text = file_bytes.decode("utf-8", errors="replace")
        lines = text.splitlines()
        pages = [{"page": i + 1, "text": line} for i, line in enumerate(lines) if line.strip()]
        return {
            "filename": filename,
            "extension": ".txt",
            "pages": pages,
            "full_text": text,
            "ocr_used": False,
            "error": None,
        }
    except Exception as e:
        return {
            "filename": filename,
            "extension": ".txt",
            "pages": [],
            "full_text": "",
            "ocr_used": False,
            "error": str(e),
        }


# cached so EasyOCR doesn't re-download the model on every call
_ocr_reader = None

def _get_ocr_reader():
    global _ocr_reader
    if _ocr_reader is None:
        import easyocr
        _ocr_reader = easyocr.Reader(["en", "de", "fr", "es", "it"], gpu=False)
    return _ocr_reader


def _tables_to_text(tables: list) -> str:
    lines = []
    for table in tables:
        for row in table:
            cleaned = [str(cell).strip() if cell else "" for cell in row]
            lines.append(" | ".join(cleaned))
    return "\n".join(lines)
