"""
Generate a varied local test dataset with three PII difficulty tiers:

  1. Obvious PII (SSN, CC, email, phone)      → caught by regex
  2. Names + addresses only                    → caught by NER
  3. Narrative / contextual PII                → caught by LLM
  4. internal_memo_* prefix                    → clean, true negatives

Files are written into the existing folder structure under TEST_DATASET_DIR.
Run: python generate_varied_data.py
"""

from __future__ import annotations

import csv
import io
import os
import random
from pathlib import Path

import fitz  # PyMuPDF
import openpyxl
from docx import Document
from docx.shared import Pt

random.seed(2025)

TEST_DATASET_DIR = Path(os.path.expanduser("~/Desktop/test_dataset"))

# ── people pool ───────────────────────────────────────────────────────────────

FIRST_NAMES = [
    "Amara", "Bjorn", "Celine", "Dmitri", "Esi", "Fatima", "Gunnar", "Hiroshi",
    "Ingrid", "Javier", "Kaito", "Leila", "Magnus", "Nadia", "Oluwaseun", "Priya",
    "Quentin", "Rania", "Soren", "Tariq", "Uma", "Viktor", "Wren", "Xiu",
    "Yusuf", "Zara", "Alistair", "Brigitte", "Callum", "Darina", "Elena", "Finn",
    "Grace", "Hugo", "Isla", "Jonas", "Kira", "Luca", "Maya", "Niall",
    "Olivia", "Patrick", "Quinn", "Rosa", "Stefan", "Tara", "Umar", "Vera",
    "William", "Yasmine",
]
LAST_NAMES = [
    "Okafor", "Lindqvist", "Dubois", "Volkov", "Mensah", "Al-Rashid", "Eriksson",
    "Tanaka", "Halvorsen", "Reyes", "Yamamoto", "Haddad", "Thorvaldsen", "Petrov",
    "Adeyemi", "Sharma", "Beaumont", "Andersson", "Kowalski", "Nakamura",
    "Fernandez", "Bergstrom", "O'Sullivan", "Papadopoulos", "Mitchell", "Clarke",
    "Bennett", "Fletcher", "Harrison", "Jensen", "Larsson", "Muller", "Nielsen",
    "Pearce", "Rhodes", "Sanderson", "Turner", "Underwood", "Vaughan", "Walsh",
]
STREETS = [
    "14 Rowanberry Close", "7 Birchwood Avenue", "23 Thornfield Lane",
    "88 Kestrel Drive", "3 Millbank Road", "51 Harrington Square",
    "120 Coppergate", "9 Foxglove Crescent", "66 Larkspur Way",
    "42 Westmoor Street", "17 Ashfield Terrace", "5 Clover Hill",
    "33 Riverside Walk", "11 Maple Court", "78 Bluebell Gardens",
]
CITIES = [
    ("Bristol", "UK"), ("Lyon", "France"), ("Malmö", "Sweden"),
    ("Gdańsk", "Poland"), ("Thessaloniki", "Greece"), ("Ghent", "Belgium"),
    ("Tampere", "Finland"), ("Braga", "Portugal"), ("Timișoara", "Romania"),
    ("Maribor", "Slovenia"), ("Novi Sad", "Serbia"), ("Aalborg", "Denmark"),
    ("Edinburgh", "UK"), ("Nantes", "France"), ("Gothenburg", "Sweden"),
]
ROLES = [
    "nurse", "secondary school teacher", "civil engineer", "accountant",
    "logistics coordinator", "social worker", "dental hygienist", "pharmacist",
    "warehouse supervisor", "graphic designer", "legal secretary", "electrician",
    "data analyst", "project manager", "HR advisor", "quantity surveyor",
]
NATIONALITIES = [
    "Nigerian", "Swedish", "French", "Polish", "Greek", "Belgian",
    "Finnish", "Portuguese", "Romanian", "Slovenian", "Serbian", "Danish",
    "Scottish", "Irish", "Dutch", "Austrian",
]
DEPARTMENTS = ["Finance", "HR", "Legal", "IT", "Operations", "Marketing", "Medical", "Sales", "Executive"]

FOLDER_MAP = {
    "Finance":    ["Finance/Expense Reports", "Finance/Invoices 2024", "Finance/Tax Documents"],
    "HR":         ["HR/Onboarding", "HR/Payroll Records", "HR/Performance Reviews"],
    "Legal":      ["Legal/Compliance", "Legal/Contracts", "Legal/NDAs"],
    "IT":         ["IT/Access Requests", "IT/Incident Logs"],
    "Operations": ["Operations/Incident Reports", "Operations/Vendor Contracts"],
    "Marketing":  ["Marketing/Campaign Data", "Marketing/Customer Lists"],
    "Medical":    ["Medical/Insurance Claims", "Medical/Patient Files"],
    "Sales":      ["Sales/Contracts", "Sales/Leads"],
    "Executive":  ["Executive/Board Reports"],
}

MONTHS = ["jan", "feb", "mar", "apr", "may", "jun",
          "jul", "aug", "sep", "oct", "nov", "dec"]
MONTH_NAMES = ["January", "February", "March", "April", "May", "June",
               "July", "August", "September", "October", "November", "December"]
YEARS = [2023, 2024, 2025]
DISTRICTS = [
    "Northside", "Riverside", "Old Town", "Westgate", "Eastfield",
    "Harbour", "Parklands", "Central", "Southwick", "Millfield",
]


def pick_folder(dept: str) -> Path:
    return TEST_DATASET_DIR / random.choice(FOLDER_MAP.get(dept, ["HR/Onboarding"]))


def random_person() -> dict:
    first = random.choice(FIRST_NAMES)
    last = random.choice(LAST_NAMES)
    city, country = random.choice(CITIES)
    return {
        "first": first, "last": last, "name": f"{first} {last}",
        "street": random.choice(STREETS), "city": city, "country": country,
        "address": f"{random.choice(STREETS)}, {city}, {country}",
        "role": random.choice(ROLES),
        "nationality": random.choice(NATIONALITIES),
        "age": random.randint(24, 67),
        "years": random.randint(1, 20),
    }


def random_ssn():
    return f"{random.randint(100,999)}-{random.randint(10,99)}-{random.randint(1000,9999)}"

def random_cc():
    return f"4{random.randint(100,999)} {random.randint(1000,9999)} {random.randint(1000,9999)} {random.randint(1000,9999)}"

def random_phone():
    return f"+{random.randint(1,49)}-{random.randint(100,999)}-555-{random.randint(1000,9999)}"

def random_email(p: dict):
    return f"{p['first'].lower()}.{p['last'].lower()}{random.randint(1,99)}@{random.choice(['example.com','corp.net','work.org','mail.eu'])}"

def random_dob():
    return f"{random.randint(1955,2000)}-{random.randint(1,12):02d}-{random.randint(1,28):02d}"

def random_ip():
    return f"{random.randint(10,203)}.{random.randint(0,255)}.{random.randint(0,255)}.{random.randint(1,254)}"

def random_passport(p: dict):
    code = p["country"][:2].upper()
    return f"{code}-{chr(random.randint(65,90))}{random.randint(10000000,99999999)}"

def ref():
    return random.randint(10000, 99999)

def month():
    i = random.randint(0, 11)
    return MONTHS[i], MONTH_NAMES[i]

def year():
    return random.choice(YEARS)


# ── realistic file names ───────────────────────────────────────────────────────

def doc_name(folder: str) -> str:
    """Return a realistic filename for the given folder path."""
    m, _ = month()
    y = year()
    n = ref()

    names = {
        "Expense Reports":    [f"expense_claim_{m}{y}.txt", f"expense_report_ref{n}.docx", f"reimbursement_form_{m}{y}.txt"],
        "Invoices 2024":      [f"invoice_{n}_{m}{y}.docx", f"supplier_invoice_{n}.txt", f"invoice_summary_{m}{y}.xlsx"],
        "Tax Documents":      [f"tax_return_{y}.docx", f"vat_filing_{m}{y}.txt", f"p60_form_{y}.docx"],
        "Onboarding":         [f"onboarding_pack_{m}{y}.docx", f"new_starter_form.txt", f"induction_checklist_{n}.docx"],
        "Payroll Records":    [f"payroll_{m}{y}.xlsx", f"salary_review_{y}.docx", f"payroll_summary_{m}{y}.txt"],
        "Performance Reviews":[f"performance_review_{y}_q{random.randint(1,4)}.docx", f"appraisal_form_{m}{y}.txt", f"annual_review_{y}.docx"],
        "Compliance":         [f"compliance_report_{m}{y}.docx", f"audit_findings_{y}.txt", f"gdpr_assessment_{y}.docx"],
        "Contracts":          [f"service_agreement_{n}.docx", f"contract_renewal_{m}{y}.txt", f"supplier_contract_{n}.docx"],
        "NDAs":               [f"nda_{n}_{m}{y}.docx", f"confidentiality_agreement_{n}.txt", f"nda_signed_{m}{y}.docx"],
        "Access Requests":    [f"access_request_{n}.txt", f"system_access_form_{m}{y}.docx", f"it_request_{n}.txt"],
        "Incident Logs":      [f"incident_report_{n}.txt", f"it_incident_{m}{y}.docx", f"security_log_{n}.txt"],
        "Incident Reports":   [f"incident_report_{n}.txt", f"ops_incident_{m}{y}.docx", f"near_miss_report_{n}.txt"],
        "Vendor Contracts":   [f"vendor_agreement_{n}.docx", f"supplier_terms_{m}{y}.txt", f"procurement_contract_{n}.docx"],
        "Campaign Data":      [f"campaign_brief_{m}{y}.docx", f"marketing_plan_{y}_q{random.randint(1,4)}.txt", f"outreach_data_{n}.xlsx"],
        "Customer Lists":     [f"customer_register_{m}{y}.xlsx", f"client_contacts_{y}.csv", f"account_list_{n}.xlsx"],
        "Insurance Claims":   [f"claim_form_{n}.txt", f"insurance_claim_{n}_{m}{y}.docx", f"medical_claim_ref{n}.txt"],
        "Patient Files":      [f"patient_intake_{n}.txt", f"referral_letter_{n}.docx", f"clinical_notes_{n}.txt"],
        "Contracts":          [f"sales_contract_{n}.docx", f"client_agreement_{n}_{m}{y}.txt", f"deal_summary_{n}.docx"],
        "Leads":              [f"lead_contact_{n}.txt", f"prospect_notes_{m}{y}.docx", f"sales_lead_{n}.txt"],
        "Board Reports":      [f"board_report_{m}{y}.docx", f"executive_summary_{y}_q{random.randint(1,4)}.txt", f"directors_briefing_{m}{y}.docx"],
    }
    folder_key = folder.split("/")[-1]
    options = names.get(folder_key, [f"document_{n}.txt"])
    return random.choice(options)


# ── TIER 1: regex-catchable PII ───────────────────────────────────────────────

def random_iban():
    country = random.choice(["GB", "DE", "FR", "NL", "ES", "IT", "PL", "SE"])
    check = random.randint(10, 99)
    bban = "".join([str(random.randint(0, 9)) for _ in range(16)])
    return f"{country}{check}{bban}"

def random_ipv6():
    groups = [f"{random.randint(0, 0xffff):04x}" for _ in range(8)]
    return ":".join(groups)

def random_drivers_license():
    letters = "".join(random.choices("ABCDEFGHJKLMNPRSTUVWXYZ", k=2))
    return f"DL-{letters}{random.randint(100000, 999999)}"

REGEX_TEMPLATES = [
    """\
EXPENSE CLAIM
Claimant:       {name}
Employee ID:    EMP-{eid}
Email:          {email}
Phone:          {phone}
Date of Birth:  {dob}
National Insurance No: {ssn}
Credit Card:    {cc}
Home Address:   {address}

Expenses:
  Travel:         £{travel}.00
  Accommodation:  £{hotel}.00
  Meals:          £{meals}.00

Authorised by Finance.
""",
    """\
PATIENT INTAKE FORM
Patient Name:   {name}
Date of Birth:  {dob}
NHS / ID No:    {ssn}
Email:          {email}
Phone:          {phone}
Address:        {address}
IP (portal):    {ip}
Passport:       {passport}

Presenting complaint: Routine annual check-up.
Insurance reference: POL-{pol}
""",
    """\
IT SYSTEM ACCESS REQUEST
Requestor:    {name}
Email:        {email}
Phone:        {phone}
Employee ID:  EMP-{eid}
IP Address:   {ip}
NI Number:    {ssn}
Access Level: Level 3 – Restricted
Systems:      Payroll, HR Database, Finance Portal
Status: Approved
""",
    """\
PAYROLL RECORD
Name:           {name}
Date of Birth:  {dob}
NI Number:      {ssn}
Email:          {email}
Phone:          {phone}
Address:        {address}
Bank Account:   {cc}
Passport No:    {passport}
Salary:         £{travel},000 per annum
Tax Code:       {eid}T
""",
    """\
BANK ACCOUNT CHANGE REQUEST
Account Holder: {name}
Date of Birth:  {dob}
Email:          {email}
Phone:          {phone}
Current IBAN:   {iban}
New IBAN:       {iban2}
Passport No:    {passport}
Request Date:   {dob}
Status: Pending verification
""",
    """\
SECURITY INCIDENT REPORT
Reporter:       {name}
Email:          {email}
Employee ID:    EMP-{eid}
Source IP:      {ip}
IPv6 Address:   {ipv6}
Timestamp:      2024-{ssn_short}T{eid}Z
Affected Users: {email2}, {email3}
Severity: High
Description: Unauthorised access attempt detected from external IP.
""",
    """\
SALES LEAD RECORD
Full Name:      {name}
Email:          {email}
Phone:          {phone}
Date of Birth:  {dob}
Address:        {address}
Driver Licence: {drivers_license}
Credit Card:    {cc}
Lead Source:    Cold outreach
Account Status: Prospect
""",
    """\
CONTRACTOR ONBOARDING FORM
Contractor:     {name}
Email:          {email}
Phone:          {phone}
SSN / Tax ID:   {ssn}
Home Address:   {address}
Bank IBAN:      {iban}
Passport:       {passport}
Start Date:     {dob}
Contract Ref:   CTR-{pol}
""",
    """\
USER ACCOUNT REGISTRATION
Username:       {username}
Full Name:      {name}
Email:          {email}
Phone:          {phone}
Date of Birth:  {dob}
IP Address:     {ip}
Account ID:     ACC-{eid}
Registration Date: {dob}
Status: Active
""",
    """\
CONFIDENTIAL PERSONNEL RECORD
Name:           {name}
Date of Birth:  {dob}
SSN:            {ssn}
Passport:       {passport}

Contact Information
-------------------
Email:          {email}
Phone:          {phone}
Address:        {address}
IP Address:     {ip}

Financial & Medical
--------------------
Credit Card:    {cc}
Medical ID:     MRN-{pol}
""",
]


def make_regex_txt(p: dict) -> str:
    p2, p3 = random_person(), random_person()
    return random.choice(REGEX_TEMPLATES).format(
        name=p["name"], email=random_email(p), phone=random_phone(),
        dob=random_dob(), ssn=random_ssn(), cc=random_cc(),
        address=p["address"], ip=random_ip(), passport=random_passport(p),
        eid=random.randint(10000, 99999),
        iban=random_iban(), iban2=random_iban(),
        ipv6=random_ipv6(),
        drivers_license=random_drivers_license(),
        username=f"{p['first'].lower()}.{p['last'].lower()}",
        email2=random_email(p2), email3=random_email(p3),
        ssn_short=f"{random.randint(1,12):02d}-{random.randint(1,28):02d}",
        travel=random.randint(20, 80),
        hotel=random.randint(80, 300),
        meals=random.randint(10, 50),
        pol=random.randint(100000, 999999),
    )


def make_regex_docx(p: dict) -> bytes:
    doc = Document()
    doc.add_heading("Confidential Staff Record", level=1)
    for label, value in [
        ("Full Name", p["name"]),
        ("Date of Birth", random_dob()),
        ("National Insurance No.", random_ssn()),
        ("Email", random_email(p)),
        ("Phone", random_phone()),
        ("Home Address", p["address"]),
        ("Credit Card / Bank Ref", random_cc()),
        ("Passport No.", random_passport(p)),
        ("IP Address", random_ip()),
    ]:
        para = doc.add_paragraph()
        run = para.add_run(f"{label}: ")
        run.bold = True
        run.font.size = Pt(11)
        para.add_run(value).font.size = Pt(11)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def make_regex_pdf(p: dict) -> bytes:
    doc_fitz = fitz.open()
    page = doc_fitz.new_page(width=595, height=842)
    page.insert_text((60, 60), make_regex_txt(p), fontsize=11)
    buf = io.BytesIO()
    doc_fitz.save(buf)
    return buf.getvalue()


# ── TIER 2: NER-only (names + addresses, no regex patterns) ──────────────────

NER_TEMPLATES = [
    """\
PERFORMANCE APPRAISAL — {year}
Employee:     {name}
Line Manager: {manager}
Department:   {dept}
Work Location: {city}, {country}
Home Address on file: {address}

Summary:
{name} has consistently met targets throughout {year}. {manager} commends
{name}'s ability to collaborate across teams and lead under pressure.
Recommended for promotion to senior {role} based in the {city} office.

Development goals:
- Shadow {colleague1} on the European accounts programme
- Work alongside {colleague2} during the Q4 client review cycle
""",
    """\
EMERGENCY CONTACT FORM
Employee Name:    {name}
Home Address:     {address}
Primary Contact:  {colleague1}
Relationship:     Spouse
Secondary Contact: {colleague2}
Relationship:     Parent

Notes: {name} has authorised {colleague1} and {colleague2} to be contacted
in the event of a workplace emergency. {manager} holds a copy of the signed
consent form.
""",
    """\
CANDIDATE SCREENING NOTES
Applicant:    {name}
Home Address: {address}
Position:     Senior {role}
Location:     {city}, {country}
Interviewed by: {manager}

Notes: {name} attended for interview on {date}. {manager} rated the
candidate highly on technical competency. {name} is currently based in
{city} and would not require relocation. Second interview with {colleague1}
to be arranged.
""",
    """\
COMPLAINT RECORD
Complainant:   {name}
Address:       {address}
Date Received: {date}
Assigned to:   {manager}

{name} contacted our {city} office on {date} regarding a service dispute.
{manager} was assigned to investigate and has been in direct correspondence
with {name}. {colleague1} was consulted for technical input.

Resolution: Pending. Follow-up due within 10 working days.
""",
    """\
MEETING MINUTES — {dept} Department
Date: {date}
Location: {city} office

Present:
  - {name} (Chair)
  - {colleague1}
  - {colleague2}
  - {colleague3}

{name} opened the meeting and welcomed {colleague3} who has joined from
the {city} branch. {colleague1} presented the quarterly figures.
{colleague2} raised concerns about the supplier based in {country}.

Actions:
  - {name} to circulate revised proposal by Friday
  - {colleague1} to arrange follow-up with {colleague2}
  - {colleague3} to complete induction with {name} next week
""",
    """\
CLIENT ACCOUNT NOTE
Client Name:     {name}
Home Address:    {address}
Account Manager: {manager}
Region:          {city}, {country}

{name} has been a client since {year}. Our account manager {manager}
visited {name} at their {city} address last quarter to review the
service agreement. {name} requested a revised quote which {manager}
will prepare by end of month.

Next contact: {manager} to call {name} directly.
Secondary contact: {colleague1}
""",
    """\
STAFF DIRECTORY ENTRY
Name:              {name}
Job Title:         Senior {role}
Office:            {city}, {country}
Home Address:      {address}
Reports to:        {manager}
Works closely with: {colleague1}, {colleague2}

Emergency contact: {colleague3} ({city} office)

{name} joined the {dept} team in {year} and relocated to {city}
from our {country} headquarters. Line manager: {manager}.
""",
    """\
REFERENCE LETTER
To whom it may concern,

I am writing in support of {name}, who has been employed in our
{dept} department since {year}. {name} is based in {city} and
reports directly to {manager}.

During their tenure, {name} has demonstrated exceptional skill as
a {role}. I have no hesitation in recommending {name} to any
prospective employer. Please direct any enquiries to {manager}
at our {city} office.

Yours sincerely,
{colleague1}
Head of {dept}
""",
    """\
VENDOR CONTACT REGISTER
Representative:  {name}
Company Address: {address}
Account Owner:   {manager}
Region:          {city}, {country}

{name} is the primary point of contact for our {city} supplier
arrangement. All purchase orders should be addressed to {name}
at {address}. Internal escalation: {manager}, copied to {colleague1}.

Backup contact: {colleague2} ({city} branch)
Contract reviewed: {year}
""",
]


def make_ner_txt(p: dict) -> str:
    people = [random_person() for _ in range(4)]
    m, mname = month()
    return random.choice(NER_TEMPLATES).format(
        name=p["name"], address=p["address"],
        city=p["city"], country=p["country"],
        role=p["role"], dept=random.choice(DEPARTMENTS),
        manager=people[0]["name"], colleague1=people[1]["name"],
        colleague2=people[2]["name"], colleague3=people[3]["name"],
        year=year(), date=f"{random.randint(1,28)} {mname} {year()}",
    )


def make_ner_docx(p: dict) -> bytes:
    people = [random_person() for _ in range(3)]
    doc = Document()
    doc.add_heading("Staff Record", level=1)
    doc.add_paragraph(f"Name: {p['name']}")
    doc.add_paragraph(f"Job Title: {p['role'].title()}")
    doc.add_paragraph(f"Office: {p['city']}, {p['country']}")
    doc.add_paragraph(f"Home address: {p['address']}")
    doc.add_paragraph(f"Line Manager: {people[0]['name']}")
    doc.add_paragraph(f"Works with: {people[1]['name']}, {people[2]['name']}")
    doc.add_paragraph(
        f"\n{p['name']} joined the {random.choice(DEPARTMENTS)} team in "
        f"{year()} and is based in {p['city']}. "
        f"Managed by {people[0]['name']}."
    )
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def make_ner_csv(people: list[dict]) -> bytes:
    buf = io.StringIO()
    writer = csv.writer(buf)
    writer.writerow(["Full Name", "Job Title", "Office", "Country", "Home Address", "Line Manager"])
    for p in people:
        mgr = random_person()
        writer.writerow([p["name"], p["role"].title(), p["city"], p["country"], p["address"], mgr["name"]])
    return buf.getvalue().encode()


def make_ner_xlsx(people: list[dict]) -> bytes:
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Staff Directory"
    ws.append(["Full Name", "Job Title", "Office", "Country", "Home Address", "Line Manager"])
    for p in people:
        mgr = random_person()
        ws.append([p["name"], p["role"].title(), p["city"], p["country"], p["address"], mgr["name"]])
    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


# ── TIER 3: LLM-only (narrative / contextual PII) ────────────────────────────

LLM_TEMPLATES = [
    """\
CASE NOTE — Social Services
Ref: CS-{ref}
Date: {date}

The individual referred to us is a {nationality} national residing in {city}
for approximately {years} years. They are in their {decade}s and work as a
{role}. The client lives with a partner and two dependent children in the
{district} area of {city}.

Prior to relocating to {country}, the client was employed in their home country.
They have expressed concerns about housing security and continuity of healthcare.

Risk level: Medium. Follow-up scheduled in four weeks.
""",
    """\
INSURANCE ASSESSOR NOTES
File: INS-{ref}

The policyholder is a {nationality} individual in their {decade}s, employed
as a {role} in the {city} area. They have held a policy with us for {years}
years without a prior claim.

The claimant reports damage to their vehicle sustained outside their home in
the {district} district. They state they were returning from a night shift
at the time of the incident.

Employment and address details have been independently verified.
Claim approved for further assessment.
""",
    """\
GP REFERRAL LETTER
To: Specialist Department, {city} General Hospital

I am referring a patient who is a {age}-year-old {nationality} {role}. The
patient has lived in {city} for {years} years and currently resides in the
{district} area. They have limited English and may require an interpreter.

The patient presents with symptoms consistent with early-stage hypertension.
Family history includes cardiovascular disease on the paternal side.

The patient has been informed and has given consent to this referral.
Please contact my surgery to arrange an appointment.
""",
    """\
HR CONFIDENTIAL NOTE
Re: Reasonable adjustment request — {dept} department

A member of our {dept} team has submitted a workplace adjustment request under
the Equality Act. The individual is a {nationality} national in their {decade}s
who has been with the organisation for {years} years, working as a {role}
from our {city} office.

The employee has disclosed a long-term condition affecting mobility and has
requested a ground-floor workstation and two days per week working from home.

HR recommendation: Approve. Review in six months.
""",
    """\
FINANCIAL ADVISOR CASE NOTES
Ref: FA-{ref}

The client is a {nationality} professional in their {decade}s currently
employed as a {role} in {city}. They relocated from their country of origin
approximately {years} years ago and are the sole earner in a household of three.

Financial objectives include early retirement at 58, funding university fees
for one child, and purchasing a property in the {district} area within five years.

Risk tolerance: Moderate. Portfolio last reviewed: {date}.
""",
    """\
SOCIAL CARE REFERRAL
Ref: SC-{ref}

Referral received for a {nationality} woman in her {decade}s living in the
{district} part of {city}. She has been resident in {country} for {years} years
and works part-time as a {role}. She lives alone and has limited local support.

Reason for referral: concerns raised by her employer regarding her wellbeing
following a period of extended absence. She has consented to this referral.

Allocated worker to make initial contact within five working days.
""",
    """\
PROBATION OFFICER REPORT
Ref: PO-{ref}
Date: {date}

The subject is a {nationality} male in his {decade}s, currently residing in
the {district} area of {city}. He has been employed as a {role} for {years}
years and has maintained stable accommodation since his release.

Compliance with reporting conditions: satisfactory. No further violations
recorded this quarter. Continued supervision recommended for one further
period.
""",
    """\
OCCUPATIONAL HEALTH REFERRAL
Ref: OH-{ref}
Date: {date}

The referral concerns a {nationality} employee in their {decade}s working as
a {role} in our {city} office for {years} years. The employee has disclosed a
musculoskeletal condition that may affect their ability to perform manual tasks.

Recommendation: phased return to work and ergonomic assessment of the workstation.
All communications should be treated as medically confidential.
""",
    """\
IMMIGRATION CASE FILE
Case Ref: IMM-{ref}
Date: {date}

The applicant is a {nationality} national currently residing in the {district}
district of {city}, {country}. They arrived {years} years ago and have worked
continuously as a {role}. The applicant is seeking indefinite leave to remain.

Supporting documents submitted: employment record, utility bills, and proof of
continuous residence. Decision expected within 90 days.
""",
]


def make_llm_txt(p: dict) -> str:
    m, mname = month()
    y = year()
    decade = (p["age"] // 10) * 10
    return random.choice(LLM_TEMPLATES).format(
        ref=ref(), date=f"{random.randint(1,28)} {mname} {y}",
        nationality=p["nationality"], city=p["city"], country=p["country"],
        years=p["years"], decade=decade, age=p["age"],
        role=p["role"], district=random.choice(DISTRICTS),
        dept=random.choice(DEPARTMENTS),
    )


# ── TIER 4: clean files ───────────────────────────────────────────────────────

CLEAN_TEMPLATES = [
    "Project update: The infrastructure migration is on track. All services have been tested in the staging environment. Go-live is planned for next quarter.",
    "Reminder: The office will be closed on the upcoming public holiday. Please submit all pending approvals before close of business on Friday.",
    "The Q2 budget review has been completed. Variance analysis shows a 3% underspend across the operations division. Full report available on the intranet.",
    "Server maintenance is scheduled for Sunday between 02:00 and 06:00 UTC. Users may experience intermittent access to internal systems during this window.",
    "The new supplier onboarding checklist has been updated. All procurement requests must include a completed risk assessment form from Q3 onwards.",
    "The annual fire safety inspection has been completed. All evacuation routes are clear. Updated assembly point maps have been posted on each floor.",
    "Reminder: mandatory cybersecurity training must be completed by all staff before the end of the month. Access the module via the learning portal.",
    "The software deployment was successful. Version 4.2.1 is now live across all environments. Release notes are documented in the internal wiki.",
    "Stock levels have been reviewed for Q2. Reorder points have been adjusted in line with updated demand forecasts provided by the sales team.",
    "The cross-departmental working group met on Tuesday. Key discussion points included process standardisation and the upcoming ERP integration.",
    "The quarterly board pack has been circulated to all non-executive directors. Papers are available in the secure document portal.",
    "Following the recent audit, three minor findings were identified. Remediation actions have been assigned to the relevant team leads with a 30-day deadline.",
    "The office relocation to the new premises is confirmed for the first week of next month. IT equipment will be moved over the weekend prior.",
    "All travel bookings above £500 now require pre-approval from a department head. The updated policy is effective from the first of next month.",
    "The data retention schedule has been reviewed and updated in line with current regulatory guidance. Archived records will be disposed of in accordance with the new policy.",
    "The vendor contract renewal has been approved for another two years. Pricing terms remain unchanged per the existing service agreement.",
    "Compliance training for all staff must be completed by end of quarter. The e-learning portal has been updated with the latest modules.",
    "System uptime for this month was 99.97%. Two minor incidents were logged, both resolved within SLA. Full report available in the operations dashboard.",
    "Budget allocation for the new fiscal year has been approved by the finance committee. Department heads will be notified of their respective allocations.",
    "The new product launch is scheduled for next month. Marketing materials are being finalised. Distribution channels have been confirmed with the logistics team.",
    "Audit findings from the internal review have been addressed. Corrective action plans are in place for all medium and high severity items.",
    "Please ensure all server configurations are updated before the maintenance window on Sunday. The deployment pipeline has been tested in the staging environment.",
    "The software update introduces improved caching mechanisms and resolves three known bugs. Rollback procedure is documented in the runbook.",
    "Inventory levels for Q2 have been updated in the ERP system. Reorder points have been adjusted based on demand forecasting models.",
    """\
DATA RETENTION POLICY — INTERNAL
Version: 3.1  |  Effective: 2024-01-01

1. Scope
This policy applies to all digital and physical records held by the organisation.

2. Retention Periods
   - Financial records: 7 years
   - HR records: 6 years after employment ends
   - Customer records: 5 years after last transaction
   - Marketing data: 2 years

3. Disposal
Records past their retention period must be securely deleted or shredded.
Disposal should be logged in the central register.

4. Review
This policy is reviewed annually by the Compliance team.
""",
    """\
MEETING AGENDA — Operations Review
Date: TBC
Location: Conference Room B

1. Welcome and apologies
2. Review of previous action items
3. Q3 KPI dashboard walkthrough
4. Supply chain update
5. IT infrastructure roadmap
6. AOB

Please confirm attendance by replying to this agenda. Papers will be
circulated 48 hours in advance.
""",
    """\
CHANGE REQUEST — IT
Ref: CR-2024-0147
Priority: Medium

Summary: Upgrade database connection pool from 50 to 100 connections to
resolve intermittent timeout errors observed during peak load.

Impact: Requires a 10-minute scheduled downtime. No user data affected.
Rollback: Revert configuration file and restart service.

Approvals required: IT Manager, Infrastructure Lead.
""",
    """\
SUPPLIER EVALUATION REPORT
Supplier: Northfield Logistics Ltd
Contract Ref: SUP-8841
Review Period: Q2 2024

Performance Metrics:
  On-time delivery:  94%
  Order accuracy:    99.1%
  Issue resolution:  Avg 1.8 days

Overall rating: Satisfactory. Contract recommended for renewal.
No significant service failures recorded this quarter.
""",
    """\
INCIDENT RESPONSE PLAYBOOK — EXCERPT
Category: Phishing Attack

Step 1: Isolate the affected workstation from the network immediately.
Step 2: Notify the IT Security team via the emergency hotline.
Step 3: Preserve logs and do not power off the machine.
Step 4: IT Security will conduct forensic triage within 4 hours.
Step 5: Communicate status to affected business unit head.

Do not attempt to remediate without guidance from the security team.
""",
    """\
RISK REGISTER — Q3 2024
Department: Operations

Risk ID  | Description                        | Likelihood | Impact | Owner
---------|-----------------------------------|------------|--------|-------
OPS-001  | Supply chain disruption            | Medium     | High   | Ops Lead
OPS-002  | ERP system outage                  | Low        | High   | IT Lead
OPS-003  | Regulatory non-compliance          | Low        | Medium | Compliance
OPS-004  | Key personnel absence              | Medium     | Medium | HR Lead

Review date: 30 September 2024. Next review: 31 December 2024.
""",
    """\
PRODUCT SPECIFICATION — INTERNAL DRAFT
Product: CloudSync Enterprise v2.0

Key features:
- Real-time data synchronisation across up to 50 nodes
- AES-256 encryption at rest and in transit
- Role-based access control with audit logging
- Support for PostgreSQL, MySQL, and SQL Server backends
- Maximum throughput: 50,000 transactions per second

Minimum hardware: 8-core CPU, 32 GB RAM, 500 GB SSD.
Recommended cloud tier: c5.2xlarge or equivalent.
""",
    """\
BOARD RESOLUTION — EXCERPT
Meeting held: Q2 Board Meeting
Resolution passed: Unanimously

The board resolved to approve the proposed capital expenditure budget of
£2.4 million for the data centre upgrade programme, subject to quarterly
progress reporting to the Audit Committee.

The board further resolved to delegate signing authority for individual
contracts below £250,000 to the Chief Operating Officer.
""",
]


def make_clean_txt() -> str:
    return random.choice(CLEAN_TEMPLATES)


# ── file writers ──────────────────────────────────────────────────────────────

def write_file(path: Path, content: bytes | str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    if isinstance(content, str):
        path.write_text(content)
    else:
        path.write_bytes(content)
    print(f"  {path.relative_to(TEST_DATASET_DIR)}")


def unique_name(folder: Path, base: str) -> Path:
    """Append a counter if the filename already exists."""
    p = folder / base
    stem, suffix = Path(base).stem, Path(base).suffix
    i = 2
    while p.exists():
        p = folder / f"{stem}_{i}{suffix}"
        i += 1
    return p


# ── main ──────────────────────────────────────────────────────────────────────

def main():
    people = [random_person() for _ in range(830)]

    print("=== Tier 1: Obvious PII (regex) ===")
    for p in people[:250]:
        dept = random.choice(DEPARTMENTS)
        folder = pick_folder(dept)
        ext = random.choice(["txt", "docx", "pdf"])
        name = doc_name(str(folder.relative_to(TEST_DATASET_DIR)))
        path = unique_name(folder, Path(name).stem + "(PII)." + ext)
        if ext == "txt":
            write_file(path, make_regex_txt(p))
        elif ext == "docx":
            write_file(path, make_regex_docx(p))
        else:
            write_file(path, make_regex_pdf(p))

    print("\n=== Tier 2: Names + addresses only (NER) ===")
    for p in people[250:600]:
        dept = random.choice(DEPARTMENTS)
        folder = pick_folder(dept)
        ext = random.choice(["txt", "txt", "docx"])  # weight toward txt
        name = doc_name(str(folder.relative_to(TEST_DATASET_DIR)))
        path = unique_name(folder, Path(name).stem + "(PII)." + ext)
        if ext == "docx":
            write_file(path, make_ner_docx(p))
        else:
            write_file(path, make_ner_txt(p))

    # Staff directory CSVs and XLSXs — names heavy
    print("\n  (Staff directories)")
    for i in range(12):
        batch = [random_person() for _ in range(25)]
        folder = pick_folder("HR")
        if i < 6:
            name = doc_name(str(folder.relative_to(TEST_DATASET_DIR)))
            path = unique_name(folder, Path(name).stem + "(PII).csv")
            write_file(path, make_ner_csv(batch))
        else:
            name = doc_name(str(folder.relative_to(TEST_DATASET_DIR)))
            path = unique_name(folder, Path(name).stem + "(PII).xlsx")
            write_file(path, make_ner_xlsx(batch))

    # Customer contact lists — names heavy
    print("\n  (Customer contact lists)")
    for i in range(8):
        batch = [random_person() for _ in range(30)]
        folder = pick_folder("Marketing")
        ext = "csv" if i % 2 == 0 else "xlsx"
        name = doc_name(str(folder.relative_to(TEST_DATASET_DIR)))
        path = unique_name(folder, Path(name).stem + "(PII)." + ext)
        if ext == "csv":
            write_file(path, make_ner_csv(batch))
        else:
            write_file(path, make_ner_xlsx(batch))

    print("\n=== Tier 3: Narrative / contextual PII (LLM) ===")
    for p in people[600:830]:
        dept = random.choice(DEPARTMENTS)
        folder = pick_folder(dept)
        name = doc_name(str(folder.relative_to(TEST_DATASET_DIR)))
        path = unique_name(folder, Path(name).stem + "(PII).txt")
        write_file(path, make_llm_txt(p))

    print("\n=== Tier 4: Clean files (noPII) ===")
    for i in range(150):
        dept = random.choice(DEPARTMENTS)
        folder = pick_folder(dept)
        path = folder / f"internal_memo_{i + 1:03d}(noPII).txt"
        write_file(path, make_clean_txt())

    print("\nDone.")
    total = sum(1 for p in TEST_DATASET_DIR.rglob("*") if p.is_file() and p.suffix)
    print(f"Total files in dataset: {total}")


if __name__ == "__main__":
    main()
